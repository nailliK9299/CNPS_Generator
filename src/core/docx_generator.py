# -*- coding: utf-8 -*-
"""
Tạo file Chứng Nhận Phát Sóng (.docx) từ template công ty.

Chiến lược: dùng python-docx thao tác trực tiếp trên file template gốc.
Template có cấu trúc cố định (tab stops, paragraph index), không dùng placeholder {{...}}.
"""
import os
import copy
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import IMAGE_WIDTH_CM


def generate_cnps(
    data: dict,
    screenshot_paths: list[str | Path],
    template_path: str | Path,
    output_path: str | Path,
) -> Path:
    """
    Tạo file CNPS .docx từ template và dữ liệu.
    
    Args:
        data: Dict chứa thông tin khoản mục:
            - ten_khach_hang: str
            - hop_dong: str
            - nhan_hang: str
            - chien_dich: str
            - noi_dung_khoan_muc: str
            - link_bai: str
            - don_vi: str
            - so_luong: str | int
            - ngay_dang: str (dd/mm/yyyy)
        screenshot_paths: Danh sách đường dẫn ảnh phân trang (theo thứ tự)
        template_path: Đường dẫn file template .docx
        output_path: Đường dẫn file .docx đầu ra
    
    Returns:
        Path đến file đã tạo
    
    Raises:
        FileNotFoundError: Nếu template không tồn tại
        ValueError: Nếu cấu trúc template không đúng mong đợi
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    
    if not template_path.exists():
        raise FileNotFoundError(f"Không tìm thấy file template: {template_path}")
    
    # Đảm bảo thư mục đích tồn tại
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document(str(template_path))
    
    # Kiểm tra cấu trúc template
    _validate_template(doc)
    
    # 1. Điền thông tin chung (header)
    _fill_header_info(doc, data)
    
    # 2. Điền bảng chi tiết lịch đăng tin
    _fill_detail_table(doc, data)
    
    # 3. Chèn ảnh chụp màn hình
    if screenshot_paths:
        _insert_screenshots(doc, screenshot_paths)
    
    # Lưu file
    doc.save(str(output_path))
    return output_path


def _validate_template(doc: Document) -> None:
    """Kiểm tra template có đúng cấu trúc mong đợi."""
    # Kiểm tra đủ paragraphs
    if len(doc.paragraphs) < 12:
        raise ValueError(
            f"Template không đúng cấu trúc: chỉ có {len(doc.paragraphs)} paragraphs "
            f"(cần ít nhất 12)"
        )
    
    # Kiểm tra các paragraph label
    expected_labels = {
        3: "Tên khách hàng:",
        4: "Hợp đồng:",
        5: "Nhãn hàng:",
        6: "Chiến dịch:",
    }
    for idx, label in expected_labels.items():
        para_text = doc.paragraphs[idx].text
        if label not in para_text:
            raise ValueError(
                f"Template sai cấu trúc: Paragraph [{idx}] = '{para_text}', "
                f"mong đợi chứa '{label}'"
            )
    
    # Kiểm tra có bảng chi tiết
    if len(doc.tables) < 2:
        raise ValueError(
            f"Template không đúng cấu trúc: chỉ có {len(doc.tables)} bảng (cần ít nhất 2)"
        )
    
    # Kiểm tra header bảng chi tiết
    table_headers = [cell.text.strip() for cell in doc.tables[1].rows[0].cells]
    expected_headers = ["Khoản mục", "Đơn vị", "Số lượng", "Ngày đăng"]
    if table_headers != expected_headers:
        raise ValueError(
            f"Template sai cấu trúc bảng: headers = {table_headers}, "
            f"mong đợi = {expected_headers}"
        )


def _fill_header_info(doc: Document, data: dict) -> None:
    """
    Điền thông tin vào P[3]→P[6].
    
    Template gốc có dạng:
        P[3] = "Tên khách hàng:\t"   (Run[0]="Tên khách hàng:", Run[1]="\t")
        P[4] = "Hợp đồng:\t\t"       (Run[0]="Hợp đồng:", Run[1..2]="\t")
        ...
    
    Cách làm: thêm 1 run mới vào cuối paragraph, copy font từ run đầu tiên.
    """
    fields = [
        (3, data.get("ten_khach_hang", "")),
        (4, data.get("hop_dong", "")),
        (5, data.get("nhan_hang", "")),
        (6, data.get("chien_dich", "")),
    ]
    
    for para_idx, value in fields:
        if not value:
            continue
        para = doc.paragraphs[para_idx]
        run = para.add_run(value)
        # Copy font properties từ run đầu tiên
        _copy_run_font(para.runs[0], run)


def _fill_detail_table(doc: Document, data: dict) -> None:
    """
    Điền dữ liệu vào row 1 (dòng dữ liệu) của Table 1 (bảng chi tiết).
    
    Cell[0] = Khoản mục: 2 paragraphs
        - Para 1: Nội dung quảng cáo (+ tabs padding)
        - Para 2: "Link: {url}"
    Cell[1] = Đơn vị
    Cell[2] = Số lượng
    Cell[3] = Ngày đăng
    """
    table = doc.tables[1]
    row = table.rows[1]
    
    # ── Cell 0: Khoản mục ──
    cell_khoan_muc = row.cells[0]
    # Xoá nội dung cũ (cell trống trong template)
    _clear_cell(cell_khoan_muc)
    
    # Paragraph 1: Nội dung khoản mục
    p1 = cell_khoan_muc.paragraphs[0]
    run1 = p1.add_run(data.get("noi_dung_khoan_muc", ""))
    run1.font.name = "Arial"
    run1.font.size = Pt(10)
    
    # Paragraph 2: Link bài
    p2 = cell_khoan_muc.add_paragraph()
    link_text = f'Link: {data.get("link_bai", "")}'
    run2 = p2.add_run(link_text)
    run2.font.name = "Arial"
    run2.font.size = Pt(10)
    
    # ── Cell 1: Đơn vị ──
    _set_cell_text(row.cells[1], data.get("don_vi", "Bài"))
    
    # ── Cell 2: Số lượng ──
    _set_cell_text(row.cells[2], str(data.get("so_luong", "1")))
    
    # ── Cell 3: Ngày đăng ──
    _set_cell_text(row.cells[3], data.get("ngay_dang", ""))


def _insert_screenshots(doc: Document, screenshot_paths: list[str | Path]) -> None:
    """
    Chèn ảnh chụp màn hình sau paragraph "ẢNH CHỤP MÀN HÌNH".
    
    Mỗi ảnh nằm trong 1 paragraph riêng. Giữa các ảnh có page break.
    Ảnh được scale vừa vùng in ngang (15.9cm).
    """
    # Tìm paragraph "ẢNH CHỤP MÀN HÌNH"
    target_para = None
    for p in doc.paragraphs:
        if "ẢNH CHỤP MÀN HÌNH" in p.text:
            target_para = p
            break
    
    if target_para is None:
        raise ValueError("Không tìm thấy mục 'ẢNH CHỤP MÀN HÌNH' trong template")
    
    # Chèn từng ảnh
    current_anchor = target_para
    for i, img_path in enumerate(screenshot_paths):
        img_path = Path(img_path)
        if not img_path.exists():
            raise FileNotFoundError(f"Không tìm thấy ảnh: {img_path}")
        
        # Thêm paragraph chứa ảnh duy nhất với page_break_before = True để ngắt trang sạch sẽ
        img_para = _insert_paragraph_after(current_anchor)
        img_para.paragraph_format.page_break_before = True
        img_para.paragraph_format.space_before = Pt(0)
        img_para.paragraph_format.space_after = Pt(0)
        img_para.paragraph_format.line_spacing = 1.0
        
        run_img = img_para.add_run()
        run_img.add_picture(str(img_path), width=Cm(IMAGE_WIDTH_CM))
        
        current_anchor = img_para


# ─── Hàm hỗ trợ nội bộ ───

def _insert_paragraph_after(paragraph) -> 'Paragraph':
    """
    Chèn 1 paragraph mới ngay SAU paragraph cho trước trong document body.
    Trả về paragraph mới.
    """
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    # Lấy paragraph object từ element
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def _copy_run_font(source_run, target_run) -> None:
    """Copy font properties từ source run sang target run."""
    target_run.font.name = source_run.font.name or "Arial"
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.bold is not None:
        target_run.font.bold = source_run.font.bold
    if source_run.font.italic is not None:
        target_run.font.italic = source_run.font.italic


def _clear_cell(cell) -> None:
    """Xoá toàn bộ nội dung trong cell (giữ lại 1 paragraph trống)."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    # Xoá paragraphs thừa (giữ paragraph đầu tiên)
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)


def _set_cell_text(cell, text: str) -> None:
    """Đặt text cho cell, giữ nguyên formatting mặc định."""
    _clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
